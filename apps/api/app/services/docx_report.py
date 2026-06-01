"""
DOCX Report Service
===================
Generates an editable Word document from analysis data using python-docx.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ANAC_BLUE = RGBColor(0x00, 0x3A, 0x70)
DARK      = RGBColor(0x1e, 0x29, 0x3b)
GREY      = RGBColor(0x64, 0x74, 0x8b)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = ANAC_BLUE
    run.font.bold = True


def _kv_row(table, key: str, value: str) -> None:
    row = table.add_row()
    key_cell = row.cells[0]
    val_cell  = row.cells[1]
    key_cell.text = key
    val_cell.text = value
    for para in key_cell.paragraphs:
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = GREY
            run.font.size = Pt(9)
    for para in val_cell.paragraphs:
        for run in para.runs:
            run.font.size = Pt(9)


def generate_docx(
    *,
    filename: str,
    created_at: datetime | None,
    total_rows: int,
    indicators: dict[str, Any] | None,
    alerts: list[dict] | None,
    ai_summary: dict[str, Any] | None,
    analysis_type: str = "ciclos",
) -> bytes:
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Title ─────────────────────────────────────────────────────────────
    title = doc.add_heading("ANAC Data Insight", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = ANAC_BLUE
        run.font.size = Pt(20)

    sub = doc.add_paragraph("Relatório Executivo de Análise")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.color.rgb = GREY
        run.font.size = Pt(11)

    doc.add_paragraph()

    # ── File info table ────────────────────────────────────────────────────
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = "Table Grid"
    info_table.columns[0].width = Inches(2)
    info_table.columns[1].width = Inches(4)

    _kv_row(info_table, "Arquivo", filename)
    _kv_row(info_table, "Data", created_at.strftime("%d/%m/%Y %H:%M") if created_at else "—")
    tipo_map = {"ciclos": "Ciclos de Fiscalização", "pdf": "Documento PDF", "generic": "Planilha Genérica"}
    _kv_row(info_table, "Tipo", tipo_map.get(analysis_type, analysis_type))
    _kv_row(info_table, "Registros", f"{total_rows:,}".replace(",", "."))

    doc.add_paragraph()
    ind = indicators or {}

    if analysis_type == "pdf":
        _heading(doc, "Informações do Documento", level=1)
        doc.add_paragraph(f"Páginas: {ind.get('pages', 0)}")
        doc.add_paragraph(f"Palavras: {ind.get('word_count', 0):,}".replace(",", "."))
        if ind.get("title"):
            doc.add_paragraph(f"Título: {ind['title']}")
        if ind.get("text_preview"):
            _heading(doc, "Prévia do Conteúdo", level=2)
            doc.add_paragraph(str(ind["text_preview"])[:1000])
    else:
        # ── Indicators ────────────────────────────────────────────────────
        _heading(doc, "Indicadores", level=1)
        metrics = [
            ("Total de atividades",    ind.get("total_atividades", 0)),
            ("Realizadas",             ind.get("realizadas", 0)),
            ("Agendadas",              ind.get("agendadas", 0)),
            ("Sem agendamento",        ind.get("sem_agendamento", 0)),
            ("Taxa de execução (%)",   ind.get("taxa_execucao", 0)),
            ("Taxa de agendamento (%)",ind.get("taxa_agendamento", 0)),
            ("Sem GIASO",              ind.get("sem_giaso", 0)),
            ("Sem PCDP",               ind.get("sem_pcdp", 0)),
            ("Sem processo",           ind.get("sem_processo", 0)),
            ("Locais indefinidos",     ind.get("locais_indefinidos", 0)),
            ("PCDP duplicada",         ind.get("pcdp_duplicada", 0)),
            ("Múltiplas PCDPs",        ind.get("multiplas_pcdps", 0)),
            ("Pendências críticas",    ind.get("pendencias_criticas", 0)),
        ]
        ind_table = doc.add_table(rows=1, cols=2)
        ind_table.style = "Table Grid"
        hdr = ind_table.rows[0].cells
        hdr[0].text = "Indicador"
        hdr[1].text = "Valor"
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = ANAC_BLUE

        for label, value in metrics:
            row = ind_table.add_row()
            row.cells[0].text = label
            row.cells[1].text = str(value)

        # ── By type ───────────────────────────────────────────────────────
        by_type = ind.get("by_type")
        if by_type:
            doc.add_paragraph()
            _heading(doc, "Resultado por Tipo de Ciclo", level=1)
            tipo_labels = {
                "CICLO_BASE": "Ciclo Base",
                "CICLO_DESEMPENHO": "Desempenho",
                "NAO_PROGRAMADA": "Não Programadas",
            }
            bt_table = doc.add_table(rows=1, cols=4)
            bt_table.style = "Table Grid"
            hdrs = bt_table.rows[0].cells
            for i, h in enumerate(["Tipo", "Total", "Realizadas", "Taxa exec."]):
                hdrs[i].text = h
                for para in hdrs[i].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = ANAC_BLUE
            for tipo, vals in by_type.items():
                r = bt_table.add_row()
                r.cells[0].text = tipo_labels.get(tipo, tipo)
                r.cells[1].text = str(vals.get("total_atividades", 0))
                r.cells[2].text = str(vals.get("realizadas", 0))
                r.cells[3].text = f"{vals.get('taxa_execucao', 0):.1f}%"

    # ── Alerts ────────────────────────────────────────────────────────────
    if alerts:
        doc.add_paragraph()
        _heading(doc, "Alertas e Pendências", level=1)
        alert_icons = {"error": "● CRÍTICO", "warning": "▲ ATENÇÃO", "info": "ℹ INFO"}
        for alert in alerts:
            prefix = alert_icons.get(alert.get("type", "info"), "")
            p = doc.add_paragraph()
            run = p.add_run(f"{prefix} {alert.get('category', '')} ({alert.get('count', 0)} ocorrência(s))")
            run.font.bold = True
            if alert.get("type") == "error":
                run.font.color.rgb = RGBColor(0xdc, 0x26, 0x26)
            elif alert.get("type") == "warning":
                run.font.color.rgb = RGBColor(0xd9, 0x77, 0x06)
            p2 = doc.add_paragraph(alert.get("message", ""))
            for run in p2.runs:
                run.font.size = Pt(9)

    # ── AI Summary ────────────────────────────────────────────────────────
    if ai_summary:
        resumo        = ai_summary.get("resumo_executivo") or ""
        achados       = ai_summary.get("principais_achados") or []
        riscos        = ai_summary.get("riscos_operacionais") or []
        recomendacoes = ai_summary.get("recomendacoes") or []
        plano         = ai_summary.get("plano_acao") or []

        if any([resumo, achados, riscos, recomendacoes]):
            doc.add_paragraph()
            _heading(doc, "Resumo Executivo — Inteligência Artificial", level=1)

        if resumo:
            doc.add_paragraph(resumo)

        def _list_section(title: str, items: list) -> None:
            if not items:
                return
            _heading(doc, title, level=2)
            for item in items:
                p = doc.add_paragraph(str(item), style="List Bullet")
                for run in p.runs:
                    run.font.size = Pt(9.5)

        _list_section("Principais Achados", achados)
        _list_section("Riscos Operacionais", riscos)
        _list_section("Recomendações", recomendacoes)

        if plano:
            _heading(doc, "Plano de Ação", level=2)
            plano_table = doc.add_table(rows=1, cols=3)
            plano_table.style = "Table Grid"
            ph = plano_table.rows[0].cells
            for i, h in enumerate(["Prioridade", "Ação", "Justificativa"]):
                ph[i].text = h
                for para in ph[i].paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = ANAC_BLUE
            for item in plano:
                r = plano_table.add_row()
                r.cells[0].text = str(item.get("prioridade", ""))
                r.cells[1].text = str(item.get("acao", ""))
                r.cells[2].text = str(item.get("justificativa", ""))

    # ── Footer note ───────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"Documento gerado automaticamente em {datetime.now().strftime('%d/%m/%Y às %H:%M')} "
        "pelo ANAC Data Insight."
    )
    for run in footer_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GREY

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
